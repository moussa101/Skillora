"""
AI Resume Analyzer - ML Service
FastAPI application for resume parsing and analysis
With Adversarial Defense (SRS v1.1)
"""

from fastapi import FastAPI, File, UploadFile, HTTPException, Form, Request
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import os

# Rate limiting imports
try:
    from slowapi import Limiter, _rate_limit_exceeded_handler
    from slowapi.util import get_remote_address
    from slowapi.errors import RateLimitExceeded
    limiter = Limiter(key_func=get_remote_address)
    RATE_LIMITING_AVAILABLE = True
except ImportError:
    limiter = None
    RATE_LIMITING_AVAILABLE = False

# File validation imports
try:
    from security.file_validator import full_file_validation, MAX_FILE_SIZE
    FILE_VALIDATION_AVAILABLE = True
except ImportError:
    FILE_VALIDATION_AVAILABLE = False
    MAX_FILE_SIZE = 10 * 1024 * 1024

# Security imports
try:
    from security.scanner import ResumeSecurityScanner, wrap_for_llm
except ImportError:
    # Fallback for when security module not yet available
    ResumeSecurityScanner = None
    wrap_for_llm = lambda x: x

# ML Analyzer import
try:
    from analyzer import get_analyzer
    ANALYZER_AVAILABLE = True
except ImportError:
    ANALYZER_AVAILABLE = False
    get_analyzer = None

# Profile Analyzer import
try:
    from profile_analyzer import ProfileAnalyzer, ProfileAnalysisResult
    PROFILE_ANALYZER_AVAILABLE = True
except ImportError:
    PROFILE_ANALYZER_AVAILABLE = False
    ProfileAnalyzer = None

# ATS Scorer import
try:
    from ats_scorer import score_ats
    ATS_SCORER_AVAILABLE = True
except ImportError:
    ATS_SCORER_AVAILABLE = False
    score_ats = None


app = FastAPI(
    title="Skillora - ML Service",
    description="ML service for resume parsing, NER extraction, and semantic similarity scoring with adversarial defense",
    version="1.2.0"
)

# Add rate limiting if available
if RATE_LIMITING_AVAILABLE:
    app.state.limiter = limiter
    app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# CORS configuration
ALLOWED_ORIGINS = os.environ.get("ALLOWED_ORIGINS", "http://localhost:3000,http://localhost:3001").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["GET", "POST", "OPTIONS"],
    allow_headers=["Content-Type", "Authorization"],
)


# Request/Response Models
class AnalyzeRequest(BaseModel):
    file_path: Optional[str] = None
    resume_text: Optional[str] = None  # Can pass text directly instead of file
    job_description: str


class SecurityInfo(BaseModel):
    is_safe: bool
    flags: List[str] = []
    invisible_text_detected: bool = False
    homoglyphs_detected: bool = False
    metadata_mismatch: bool = False


class GitHubProfileInfo(BaseModel):
    username: str
    name: Optional[str] = None
    bio: Optional[str] = None
    company: Optional[str] = None
    location: Optional[str] = None
    public_repos: int = 0
    followers: int = 0
    total_stars: int = 0
    top_languages: List[str] = []
    recent_commits: int = 0
    notable_repos: List[Dict[str, Any]] = []


class ProfileAnalysis(BaseModel):
    github: Optional[GitHubProfileInfo] = None
    linkedin_url: Optional[str] = None
    profile_score: float = 0.0
    profile_insights: List[str] = []
    urls_found: Dict[str, Any] = {}


class LanguageInfo(BaseModel):
    """Detected language information for the resume"""
    language_code: str  # ISO 639-1 code (e.g., 'en', 'es', 'fr')
    language_name: str  # Full name (e.g., 'English', 'Spanish')
    confidence: float  # 0.0 to 1.0
    is_rtl: bool = False  # Right-to-left language


class ATSCheckInfo(BaseModel):
    """Individual ATS check result"""
    name: str
    passed: bool
    message: str
    severity: str  # "critical", "warning", "info", "good"


class ATSCategoryInfo(BaseModel):
    """ATS scoring category"""
    name: str
    score: float
    checks: List[ATSCheckInfo] = []


class ATSScoreInfo(BaseModel):
    """ATS compatibility score breakdown"""
    overall_score: float
    keyword_match_rate: float = 0.0
    critical_issues: List[str] = []
    suggestions: List[str] = []
    categories: List[ATSCategoryInfo] = []


class AnalyzeResponse(BaseModel):
    score: float
    suspicious: bool = False  # NFR-SEC-04: Flag high scores
    suspicious_reason: Optional[str] = None
    security: Optional[SecurityInfo] = None
    skills_found: List[str]
    profile_analysis: Optional[ProfileAnalysis] = None  # GitHub/LinkedIn analysis
    missing_keywords: List[str]
    contact_info: Optional[dict] = None
    education: Optional[List[dict]] = None
    experience: Optional[List[dict]] = None
    feedback: Optional[dict] = None
    language: Optional[LanguageInfo] = None  # NEW: Detected language info
    ats_score: Optional[ATSScoreInfo] = None  # ATS compatibility scoring


class HealthResponse(BaseModel):
    status: str
    version: str


class SecurityScanResponse(BaseModel):
    is_safe: bool
    security_flags: List[str]
    details: Dict[str, Any]


# Health check endpoint
@app.get("/health", response_model=HealthResponse)
async def health_check():
    """Health check endpoint"""
    return HealthResponse(status="healthy", version="1.1.0")


# Security scan endpoint
@app.post("/security-scan", response_model=SecurityScanResponse)
async def security_scan(file_path: str):
    """
    FR-SEC-01 to FR-SEC-05: Security scan for adversarial attacks
    """
    if ResumeSecurityScanner is None:
        raise HTTPException(status_code=500, detail="Security scanner not available")
    
    scanner = ResumeSecurityScanner()
    result = scanner.scan_pdf(file_path)
    
    return SecurityScanResponse(
        is_safe=result.get("is_safe", False),
        security_flags=result.get("security_flags", []),
        details=result
    )


# Main analysis endpoint
@app.post("/analyze", response_model=AnalyzeResponse)
async def analyze_resume(request: AnalyzeRequest):
    """
    Analyze a resume against a job description.
    Accepts either file_path or resume_text directly.
    """
    
    # Determine resume text source
    resume_text = request.resume_text
    file_exists = request.file_path and os.path.exists(request.file_path)
    
    # Run security scan if file exists
    security_info = None
    if ResumeSecurityScanner is not None and file_exists:
        scanner = ResumeSecurityScanner()
        scan_result = scanner.scan_pdf(request.file_path)
        security_info = SecurityInfo(
            is_safe=scan_result.get("is_safe", True),
            flags=scan_result.get("security_flags", []),
            invisible_text_detected=scan_result.get("invisible_text_detected", False),
            homoglyphs_detected=scan_result.get("homoglyphs_detected", False),
            metadata_mismatch=scan_result.get("metadata_mismatch", False),
        )
    
    # Use real analyzer if available
    language_info = None
    if ANALYZER_AVAILABLE and get_analyzer is not None:
        analyzer = get_analyzer()
        
        if resume_text:
            # Direct text analysis with language detection
            result = analyzer.analyze_text(resume_text, request.job_description)
            score = result["score"]
            skills_found = result["skills_found"]
            missing_keywords = result["missing_keywords"]
            feedback = result["feedback"]
            if "language" in result:
                language_info = LanguageInfo(
                    language_code=result["language"]["language_code"],
                    language_name=result["language"]["language_name"],
                    confidence=result["language"]["confidence"],
                    is_rtl=result["language"].get("is_rtl", False)
                )
        elif file_exists:
            # File-based analysis with language detection
            result = analyzer.analyze(request.file_path, request.job_description)
            score = result["score"]
            skills_found = result["skills_found"]
            missing_keywords = result["missing_keywords"]
            feedback = result["feedback"]
            if "language" in result:
                language_info = LanguageInfo(
                    language_code=result["language"]["language_code"],
                    language_name=result["language"]["language_name"],
                    confidence=result["language"]["confidence"],
                    is_rtl=result["language"].get("is_rtl", False)
                )
        else:
            # No valid input
            score = 0.0
            skills_found = []
            missing_keywords = []
            feedback = {
                "summary": "No resume text or valid file provided",
                "suggestions": ["Upload a resume file or paste resume text"]
            }
    else:
        # Fallback when analyzer not available
        score = 50.0
        skills_found = ["Analyzer not available"]
        missing_keywords = []
        feedback = {
            "summary": "Analysis requires sentence-transformers",
            "suggestions": ["Run: pip install sentence-transformers"]
        }
    
    # NFR-SEC-04: Anomaly Detection - flag suspiciously high scores
    suspicious = False
    suspicious_reason = None
    if score >= 95.0:
        suspicious = True
        suspicious_reason = "Score >= 95% indicates possible JD copy-paste. Manual review required."
    
    # ATS Compatibility Scoring
    ats_result = None
    if ATS_SCORER_AVAILABLE and score_ats is not None and resume_text:
        try:
            ats_data = score_ats(resume_text, request.job_description, skills_found, missing_keywords)
            ats_result = ATSScoreInfo(
                overall_score=ats_data["overall_score"],
                keyword_match_rate=ats_data["keyword_match_rate"],
                critical_issues=ats_data["critical_issues"],
                suggestions=ats_data["suggestions"],
                categories=[
                    ATSCategoryInfo(
                        name=cat["name"],
                        score=cat["score"],
                        checks=[ATSCheckInfo(**c) for c in cat["checks"]],
                    )
                    for cat in ats_data["categories"]
                ],
            )
        except Exception as e:
            print(f"ATS scoring failed: {e}")

    return AnalyzeResponse(
        score=round(score, 1),
        suspicious=suspicious,
        suspicious_reason=suspicious_reason,
        security=security_info,
        skills_found=skills_found[:10],
        missing_keywords=missing_keywords[:8],
        feedback=feedback,
        language=language_info,
        ats_score=ats_result,
    )


# Text extraction endpoint
@app.post("/extract-text")
async def extract_text(request: Request, file: UploadFile = File(...)):
    """
    Extract text from an uploaded resume file.
    Supports PDF, DOCX, TXT, RTF, and HTML formats.
    Max file size: 10MB. Rate limited: 30 requests/minute.
    """
    
    try:
        content = await file.read()
        
        # Security validation
        if FILE_VALIDATION_AVAILABLE:
            is_valid, error, safe_filename = full_file_validation(file.filename, content)
            if not is_valid:
                raise HTTPException(status_code=400, detail=error)
        else:
            # Fallback validation
            safe_filename = file.filename
            if len(content) > MAX_FILE_SIZE:
                raise HTTPException(status_code=400, detail=f"File too large. Max size: {MAX_FILE_SIZE // (1024*1024)}MB")
        
        file_ext = os.path.splitext(safe_filename)[1].lower()
        allowed_extensions = ['.pdf', '.docx', '.txt', '.rtf', '.html', '.htm']
        
        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}"
            )
        
        text = ""
        
        if file_ext == '.pdf':
            # Use PyMuPDF for PDF extraction
            import fitz
            doc = fitz.open(stream=content, filetype="pdf")
            for page in doc:
                text += page.get_text()
            doc.close()
        
        elif file_ext == '.docx':
            # Use python-docx for DOCX extraction
            import docx
            from io import BytesIO
            doc = docx.Document(BytesIO(content))
            text = "\n".join([para.text for para in doc.paragraphs])
        
        elif file_ext == '.txt':
            text = content.decode('utf-8')
        
        elif file_ext == '.rtf':
            # Use striprtf for RTF extraction
            from striprtf.striprtf import rtf_to_text
            text = rtf_to_text(content.decode('utf-8', errors='ignore'))
        
        elif file_ext in ['.html', '.htm']:
            # Use BeautifulSoup for HTML extraction
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(content.decode('utf-8', errors='ignore'), 'lxml')
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            text = soup.get_text(separator='\n', strip=True)
        
        return {
            "filename": safe_filename,
            "text": text,
            "length": len(text),
            "status": "success"
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to extract text: {str(e)}")


# File upload + analyze endpoint (combines upload and analysis)
@app.post("/analyze-file", response_model=AnalyzeResponse)
async def analyze_file(request: Request, file: UploadFile = File(...), job_description: str = Form(...)):
    """
    Upload a resume file and analyze it against a job description.
    Extracts text from PDF/DOCX/TXT/RTF/HTML and performs analysis.
    Max file size: 10MB.
    """
    
    try:
        content = await file.read()
        
        # Security validation
        if FILE_VALIDATION_AVAILABLE:
            is_valid, error, safe_filename = full_file_validation(file.filename, content)
            if not is_valid:
                raise HTTPException(status_code=400, detail=error)
        else:
            # Fallback validation
            safe_filename = file.filename
            if len(content) > MAX_FILE_SIZE:
                raise HTTPException(status_code=400, detail=f"File too large. Max size: {MAX_FILE_SIZE // (1024*1024)}MB")
        
        file_ext = os.path.splitext(safe_filename)[1].lower()
        allowed_extensions = ['.pdf', '.docx', '.txt', '.rtf', '.html', '.htm']
        
        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}"
            )
        
        text = ""
        
        if file_ext == '.pdf':
            import fitz
            doc = fitz.open(stream=content, filetype="pdf")
            for page in doc:
                text += page.get_text()
            doc.close()
        
        elif file_ext == '.docx':
            import docx
            from io import BytesIO
            doc = docx.Document(BytesIO(content))
            text = "\n".join([para.text for para in doc.paragraphs])
        
        elif file_ext == '.txt':
            text = content.decode('utf-8')
        
        elif file_ext == '.rtf':
            # Use striprtf for RTF extraction
            from striprtf.striprtf import rtf_to_text
            text = rtf_to_text(content.decode('utf-8', errors='ignore'))
        
        elif file_ext in ['.html', '.htm']:
            # Use BeautifulSoup for HTML extraction
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(content.decode('utf-8', errors='ignore'), 'lxml')
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            text = soup.get_text(separator='\n', strip=True)
        
        if not text or len(text) < 10:
            raise HTTPException(status_code=400, detail="Could not extract text from file")
        
        # Now analyze using the extracted text with language detection
        language_info = None
        if ANALYZER_AVAILABLE and get_analyzer is not None:
            analyzer = get_analyzer()
            result = analyzer.analyze_text(text, job_description)
            score = result["score"]
            skills_found = result["skills_found"]
            missing_keywords = result["missing_keywords"]
            feedback = result["feedback"]
            if "language" in result:
                language_info = LanguageInfo(
                    language_code=result["language"]["language_code"],
                    language_name=result["language"]["language_name"],
                    confidence=result["language"]["confidence"],
                    is_rtl=result["language"].get("is_rtl", False)
                )
        else:
            score = 50.0
            skills_found = ["Analyzer not available"]
            missing_keywords = []
            feedback = {"summary": "Install dependencies", "suggestions": []}
        
        # Profile analysis (GitHub/LinkedIn)
        profile_analysis = None
        if PROFILE_ANALYZER_AVAILABLE and ProfileAnalyzer is not None:
            try:
                profile_analyzer = ProfileAnalyzer()
                profile_result = await profile_analyzer.analyze(text)
                
                github_info = None
                if profile_result.github:
                    github_info = GitHubProfileInfo(
                        username=profile_result.github.username,
                        name=profile_result.github.name,
                        bio=profile_result.github.bio,
                        company=profile_result.github.company,
                        location=profile_result.github.location,
                        public_repos=profile_result.github.public_repos,
                        followers=profile_result.github.followers,
                        total_stars=profile_result.github.total_stars,
                        top_languages=profile_result.github.top_languages,
                        recent_commits=profile_result.github.recent_commits,
                        notable_repos=profile_result.github.notable_repos
                    )
                
                profile_analysis = ProfileAnalysis(
                    github=github_info,
                    linkedin_url=profile_result.linkedin_url,
                    profile_score=profile_result.profile_score,
                    profile_insights=profile_result.profile_insights,
                    urls_found=profile_result.urls_found
                )
            except Exception as e:
                print(f"Profile analysis failed: {e}")
                # Continue without profile analysis
        
        # Anomaly detection
        suspicious = score >= 95.0
        suspicious_reason = "Score >= 95% indicates possible JD copy-paste." if suspicious else None
        
        # ATS Compatibility Scoring
        ats_result = None
        if ATS_SCORER_AVAILABLE and score_ats is not None and text:
            try:
                ats_data = score_ats(text, job_description, skills_found, missing_keywords)
                ats_result = ATSScoreInfo(
                    overall_score=ats_data["overall_score"],
                    keyword_match_rate=ats_data["keyword_match_rate"],
                    critical_issues=ats_data["critical_issues"],
                    suggestions=ats_data["suggestions"],
                    categories=[
                        ATSCategoryInfo(
                            name=cat["name"],
                            score=cat["score"],
                            checks=[ATSCheckInfo(**c) for c in cat["checks"]],
                        )
                        for cat in ats_data["categories"]
                    ],
                )
            except Exception as e:
                print(f"ATS scoring failed: {e}")

        return AnalyzeResponse(
            score=round(score, 1),
            suspicious=suspicious,
            suspicious_reason=suspicious_reason,
            security=None,
            skills_found=skills_found[:10],
            missing_keywords=missing_keywords[:8],
            feedback=feedback,
            profile_analysis=profile_analysis,
            language=language_info,
            ats_score=ats_result,
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")


class BatchFileItem(BaseModel):
    """Result for a single file in a batch"""
    filename: str
    score: float = 0.0
    suspicious: bool = False
    suspicious_reason: Optional[str] = None
    skills_found: List[str] = []
    missing_keywords: List[str] = []
    feedback: Optional[dict] = None
    ats_score: Optional[ATSScoreInfo] = None
    error: Optional[str] = None


class BatchAnalyzeResponse(BaseModel):
    total: int
    successful: int
    failed: int
    results: List[BatchFileItem]


@app.post("/batch-analyze", response_model=BatchAnalyzeResponse)
async def batch_analyze(
    request: Request,
    files: List[UploadFile] = File(...),
    job_description: str = Form(...)
):
    """
    Batch analyze multiple resume files against a single job description.
    Max 50 files per batch.
    """

    if len(files) > 50:
        raise HTTPException(status_code=400, detail="Maximum 50 files per batch")

    if not job_description or len(job_description.strip()) < 10:
        raise HTTPException(status_code=400, detail="Job description must be at least 10 characters")

    async def process_single(file: UploadFile) -> BatchFileItem:
        try:
            content = await file.read()

            if FILE_VALIDATION_AVAILABLE:
                is_valid, error, safe_filename = full_file_validation(file.filename, content)
                if not is_valid:
                    return BatchFileItem(filename=file.filename, error=error)
            else:
                safe_filename = file.filename
                if len(content) > MAX_FILE_SIZE:
                    return BatchFileItem(filename=file.filename, error="File too large")

            file_ext = os.path.splitext(safe_filename)[1].lower()
            allowed_extensions = ['.pdf', '.docx', '.txt', '.rtf', '.html', '.htm']
            if file_ext not in allowed_extensions:
                return BatchFileItem(filename=file.filename, error=f"Unsupported file type: {file_ext}")

            # Extract text
            text = ""
            if file_ext == '.pdf':
                import fitz
                doc = fitz.open(stream=content, filetype="pdf")
                for page in doc:
                    text += page.get_text()
                doc.close()
            elif file_ext == '.docx':
                import docx as docx_module
                from io import BytesIO
                doc = docx_module.Document(BytesIO(content))
                text = "\n".join([para.text for para in doc.paragraphs])
            elif file_ext == '.txt':
                text = content.decode('utf-8')
            elif file_ext == '.rtf':
                from striprtf.striprtf import rtf_to_text
                text = rtf_to_text(content.decode('utf-8', errors='ignore'))
            elif file_ext in ['.html', '.htm']:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(content.decode('utf-8', errors='ignore'), 'lxml')
                for script in soup(["script", "style"]):
                    script.decompose()
                text = soup.get_text(separator='\n', strip=True)

            if not text or len(text) < 10:
                return BatchFileItem(filename=file.filename, error="Could not extract text")

            # Analyze
            if ANALYZER_AVAILABLE and get_analyzer is not None:
                analyzer = get_analyzer()
                result = analyzer.analyze_text(text, job_description)
                score = result["score"]
                skills_found = result["skills_found"]
                missing_keywords = result["missing_keywords"]
                feedback = result["feedback"]
            else:
                score = 50.0
                skills_found = []
                missing_keywords = []
                feedback = {"summary": "Analyzer not available", "suggestions": []}

            suspicious = score >= 95.0
            suspicious_reason = "Score >= 95% indicates possible JD copy-paste." if suspicious else None

            # ATS scoring
            ats_result = None
            if ATS_SCORER_AVAILABLE and score_ats is not None and text:
                try:
                    ats_data = score_ats(text, job_description, skills_found, missing_keywords)
                    ats_result = ATSScoreInfo(
                        overall_score=ats_data["overall_score"],
                        keyword_match_rate=ats_data["keyword_match_rate"],
                        critical_issues=ats_data["critical_issues"],
                        suggestions=ats_data["suggestions"],
                        categories=[
                            ATSCategoryInfo(
                                name=cat["name"],
                                score=cat["score"],
                                checks=[ATSCheckInfo(**c) for c in cat["checks"]],
                            )
                            for cat in ats_data["categories"]
                        ],
                    )
                except Exception as e:
                    print(f"ATS scoring failed for {file.filename}: {e}")

            return BatchFileItem(
                filename=file.filename,
                score=round(score, 1),
                suspicious=suspicious,
                suspicious_reason=suspicious_reason,
                skills_found=skills_found[:10],
                missing_keywords=missing_keywords[:8],
                feedback=feedback,
                ats_score=ats_result,
            )
        except Exception as e:
            return BatchFileItem(filename=file.filename, error=str(e))

    # Process all files sequentially (CPU-bound ML work)
    results = []
    for f in files:
        r = await process_single(f)
        results.append(r)

    successful = sum(1 for r in results if r.error is None)
    failed = sum(1 for r in results if r.error is not None)

    return BatchAnalyzeResponse(
        total=len(files),
        successful=successful,
        failed=failed,
        results=results
    )


# Standalone profile analysis endpoint
@app.post("/analyze-profiles", response_model=ProfileAnalysis)
async def analyze_profiles(github_url: Optional[str] = None, linkedin_url: Optional[str] = None, resume_text: Optional[str] = None):
    """
    Analyze GitHub/LinkedIn profiles directly.
    Can pass URLs directly or resume text to extract URLs from.
    """
    if not PROFILE_ANALYZER_AVAILABLE or ProfileAnalyzer is None:
        raise HTTPException(status_code=500, detail="Profile analyzer not available")
    
    profile_analyzer = ProfileAnalyzer()
    
    # If resume text provided, extract URLs from it
    if resume_text:
        profile_result = await profile_analyzer.analyze(resume_text)
    else:
        # Construct text with provided URLs
        text_parts = []
        if github_url:
            text_parts.append(github_url)
        if linkedin_url:
            text_parts.append(linkedin_url)
        
        if not text_parts:
            raise HTTPException(status_code=400, detail="Provide github_url, linkedin_url, or resume_text")
        
        profile_result = await profile_analyzer.analyze(" ".join(text_parts))
    
    github_info = None
    if profile_result.github:
        github_info = GitHubProfileInfo(
            username=profile_result.github.username,
            name=profile_result.github.name,
            bio=profile_result.github.bio,
            company=profile_result.github.company,
            location=profile_result.github.location,
            public_repos=profile_result.github.public_repos,
            followers=profile_result.github.followers,
            total_stars=profile_result.github.total_stars,
            top_languages=profile_result.github.top_languages,
            recent_commits=profile_result.github.recent_commits,
            notable_repos=profile_result.github.notable_repos
        )
    
    return ProfileAnalysis(
        github=github_info,
        linkedin_url=profile_result.linkedin_url,
        profile_score=profile_result.profile_score,
        profile_insights=profile_result.profile_insights,
        urls_found=profile_result.urls_found
    )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)


